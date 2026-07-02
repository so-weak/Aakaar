"""cap.word_read — extract text, paragraphs and tables from a .docx.

Reads a Word document (``aakaar://`` URI) with python-docx and returns the
full text, the list of paragraphs, and each table as a 2-D array of cell
strings. Read-only (``side_effecting=False``). python-docx is imported lazily.
"""

from __future__ import annotations

import io
import logging
from typing import Any

from pydantic import BaseModel, ConfigDict, Field

from aakaar_caps.context import CapabilityContext
from aakaar_caps.spec import CapabilitySpec

logger = logging.getLogger(__name__)
CAP_REF = "cap.word_read"


class _Inputs(BaseModel):
    model_config = ConfigDict(extra="forbid")
    file_uri: str = Field(description="Managed-storage URI (aakaar://...) of the .docx document to read.")


class _Outputs(BaseModel):
    text: str = Field(description="Full document text (paragraphs joined by newlines).")
    paragraphs: list[str] = Field(description="Non-empty paragraph texts, in document order.")
    tables: list[list[list[str]]] = Field(
        description="Each table as a 2-D array of cell strings (row-major).",
    )


SPEC = CapabilitySpec(
    ref=CAP_REF,
    description=(
        "Read a .docx document from managed storage with python-docx and return the full text, "
        "the list of paragraphs (in order), and each table as a 2-D array of cell strings. "
        "Offline, read-only."
    ),
    input_schema=_Inputs,
    output_schema=_Outputs,
    secrets=(),
    tags=("office", "word"),
    side_effecting=False,
)


def _require_docx() -> Any:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover - exercised only when dep absent
        raise RuntimeError(
            "cap.word_read needs python-docx — install aakaar-capabilities[automation]"
        ) from exc
    return docx


async def run(ctx: CapabilityContext, inputs: dict[str, Any]) -> dict[str, Any]:
    docx = _require_docx()
    file_uri = inputs["file_uri"]

    data = await ctx.read_object(file_uri)
    logger.info("cap.word_read start run_id=%s uri=%s", ctx.run_id, file_uri)
    doc = docx.Document(io.BytesIO(data))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(p.text for p in doc.paragraphs)
    tables: list[list[list[str]]] = [
        [[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables
    ]

    logger.info("cap.word_read ok run_id=%s paras=%d tables=%d", ctx.run_id, len(paragraphs), len(tables))
    return {"text": text, "paragraphs": paragraphs, "tables": tables}
