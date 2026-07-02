"""cap.word_write — build a .docx from paragraphs (and optional headings) and store it.

Takes a list of ``paragraphs`` and optional ``headings`` (interleaved before the
paragraphs, or supplied inline as ``{"heading": text, "level": n}`` items),
builds a document with python-docx, writes it to managed storage via
``ctx.write_object`` and returns the resulting ``aakaar://`` URI. Produces an
artifact, so ``side_effecting=True``. python-docx is imported lazily.
"""

from __future__ import annotations

import io
import logging
from typing import Any

from pydantic import BaseModel, ConfigDict, Field

from aakaar_caps.context import CapabilityContext
from aakaar_caps.spec import CapabilitySpec

logger = logging.getLogger(__name__)
CAP_REF = "cap.word_write"


class _Inputs(BaseModel):
    model_config = ConfigDict(extra="forbid")
    paragraphs: list[Any] = Field(
        description=(
            "Document body. Each item is either a plain string (a normal paragraph) or a dict "
            '{"heading": "...", "level": 1} to emit a heading at the given level (1-9).'
        ),
    )
    title: str | None = Field(
        default=None,
        description="Optional document title, added as a level-0 heading at the top.",
    )
    filename: str = Field(default="document.docx", description="Base filename for the stored artifact.")


class _Outputs(BaseModel):
    document_uri: str = Field(description="Managed-storage URI (aakaar://...) of the produced .docx document.")
    paragraph_count: int = Field(description="Number of body items written.")


SPEC = CapabilitySpec(
    ref=CAP_REF,
    description=(
        "Build a .docx document from a list of paragraphs (plain strings, or "
        '{"heading","level"} dicts for headings) plus an optional title, with python-docx, '
        "store it in managed storage, and return the artifact URI. Offline."
    ),
    input_schema=_Inputs,
    output_schema=_Outputs,
    secrets=(),
    tags=("office", "word"),
    side_effecting=True,
)


def _require_docx() -> Any:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover - exercised only when dep absent
        raise RuntimeError(
            "cap.word_write needs python-docx — install aakaar-capabilities[automation]"
        ) from exc
    return docx


async def run(ctx: CapabilityContext, inputs: dict[str, Any]) -> dict[str, Any]:
    docx = _require_docx()
    paragraphs = inputs["paragraphs"]
    title = inputs.get("title")
    filename = str(inputs.get("filename", "document.docx"))

    logger.info("cap.word_write start run_id=%s items=%d", ctx.run_id, len(paragraphs))
    doc = docx.Document()
    if title:
        doc.add_heading(str(title), level=0)

    for item in paragraphs:
        if isinstance(item, dict) and "heading" in item:
            level = int(item.get("level", 1))
            level = max(1, min(9, level))
            doc.add_heading(str(item["heading"]), level=level)
        else:
            doc.add_paragraph(str(item))

    buf = io.BytesIO()
    doc.save(buf)
    uri = await ctx.write_object(filename, buf.getvalue())
    logger.info("cap.word_write ok run_id=%s uri=%s items=%d", ctx.run_id, uri, len(paragraphs))
    return {"document_uri": uri, "paragraph_count": len(paragraphs)}
